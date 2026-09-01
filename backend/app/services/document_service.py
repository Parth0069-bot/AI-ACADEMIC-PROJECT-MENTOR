"""
Document generation and persistent storage.

Generates:
- Synopsis
- Methodology
- Progress Report

The generated DOCX is:
1. Created using Gemini + python-docx.
2. Uploaded to the private Supabase Storage bucket
   "generated-documents".
3. Recorded in the public.documents table.
4. Returned by the API so the browser can download it immediately.

This means documents remain available even if the student's
downloaded copy is deleted from their laptop.
"""

import json
import os
import re
import tempfile
import uuid
from datetime import datetime, timezone

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from google.genai import types

from app.core.gemini_client import get_gemini_client
from app.core.config import settings
from app.core.idea_repository import fetch_idea_with_student_context
from app.core.feedback_repository import fetch_latest_agent_feedback
from app.core.mentor_repository import fetch_checkins
from app.core.supabase_client import get_supabase
from app.schemas.document import DocumentType, GeneratedDocumentContent


AGENT_NAMES = [
    "feasibility_agent",
    "scope_agent",
    "technology_agent",
    "timeline_agent",
    "risk_agent",
]

DOCUMENT_BUCKET = "generated-documents"

DOCX_MEDIA_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
)


DOCUMENT_INSTRUCTIONS = {
    DocumentType.synopsis: (
        "Write a formal project SYNOPSIS suitable for faculty review. Include sections such as "
        "Introduction, Problem Statement, Objectives, Proposed Solution, Scope, and Expected "
        "Outcome. Base every claim strictly on the provided project data -- do not invent details."
    ),

    DocumentType.methodology: (
        "Write a formal METHODOLOGY document suitable for faculty review. Include sections such as "
        "Overview, Technology Stack & Justification, System Architecture / Approach, Implementation "
        "Plan (organized by phase or week if timeline data is available), and Evaluation Plan. Base "
        "every claim strictly on the provided project data -- do not invent details."
    ),

    DocumentType.progress_report: (
        "Write a formal PROGRESS REPORT suitable for faculty review, covering work completed to "
        "date. Include sections such as Summary, Work Completed So Far (organized by week using "
        "the check-in history), Current Risks & Blockers, Timeline Adjustments Made, and Next "
        "Steps. Base every claim strictly on the provided check-in and agent data -- do not invent "
        "or exaggerate progress that isn't reflected in the check-in history."
    ),
}


DOCUMENT_SYSTEM_PROMPT = """You are the Documentation Agent inside an AI academic project mentor \
platform. You generate formal, faculty-facing documents on demand, grounded strictly in the \
project data you're given. Write in a professional, academic tone. Do not use markdown formatting \
symbols (no #, *, **) -- section bodies should be plain prose paragraphs, since this will be \
rendered directly into a Word document.

Respond with a single JSON object matching exactly this shape:
{
  "title": "<document title, including the project title>",
  "subtitle": "<short subtitle, e.g. 'Progress Report -- Week 6', or empty string>",
  "sections": [
    {
      "heading": "<section heading>",
      "body": "<one or more full paragraphs of prose>"
    }
  ]
}"""


def _extract_json(raw_text: str) -> dict:
    """Extract valid JSON from Gemini's response."""

    cleaned = raw_text.strip()

    fenced_match = re.match(
        r"^```(?:json)?\s*(.*?)\s*```$",
        cleaned,
        re.DOTALL,
    )

    if fenced_match:
        cleaned = fenced_match.group(1)

    try:
        return json.loads(cleaned)

    except json.JSONDecodeError as exc:
        raise ValueError(
            "Document agent did not return valid JSON. "
            f"Raw response was:\n{raw_text[:500]}"
        ) from exc


def _gather_context(idea_id: str) -> dict:
    """
    Collect all project information needed by the Documentation Agent.
    """

    context = fetch_idea_with_student_context(idea_id)

    agent_feedback = {}

    for name in AGENT_NAMES:
        fb = fetch_latest_agent_feedback(idea_id, name)

        if fb and fb.details:
            agent_feedback[name] = fb.details

    checkins = fetch_checkins(idea_id)

    return {
        "idea": context.idea.model_dump(mode="json"),
        "skills": [s.model_dump() for s in context.skills],
        "agent_feedback": agent_feedback,
        "checkin_history": checkins,
    }


def generate_document_content(
    idea_id: str,
    document_type: DocumentType,
) -> GeneratedDocumentContent:
    """
    Uses Gemini to generate structured academic document content.
    """

    client = get_gemini_client()

    context = _gather_context(idea_id)

    user_prompt = (
        f"{DOCUMENT_INSTRUCTIONS[document_type]}\n\n"
        f"Here is the full project data:\n\n"
        f"{json.dumps(context, indent=2, default=str)}"
    )

    response = client.models.generate_content(
        model=settings.gemini_model,
        contents=user_prompt,
        config=types.GenerateContentConfig(
            system_instruction=DOCUMENT_SYSTEM_PROMPT,
            response_mime_type="application/json",
            max_output_tokens=8000,
        ),
    )

    parsed = _extract_json(response.text)

    return GeneratedDocumentContent(**parsed)


def render_docx(
    content: GeneratedDocumentContent,
    document_type: DocumentType,
) -> str:
    """
    Creates a DOCX file in the temporary directory.

    The file is later uploaded to Supabase Storage by
    store_generated_document().
    """

    doc = Document()

    # ------------------------------------------------------------
    # Title
    # ------------------------------------------------------------

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_run = title_p.add_run(content.title)
    title_run.bold = True
    title_run.font.size = Pt(20)

    # ------------------------------------------------------------
    # Subtitle
    # ------------------------------------------------------------

    if content.subtitle:
        subtitle_p = doc.add_paragraph()
        subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle_run = subtitle_p.add_run(content.subtitle)
        subtitle_run.italic = True
        subtitle_run.font.size = Pt(13)

    # ------------------------------------------------------------
    # Generated date
    # ------------------------------------------------------------

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    date_run = date_p.add_run(
        datetime.now().strftime("Generated on %B %d, %Y")
    )

    date_run.font.size = Pt(9)
    date_run.italic = True

    doc.add_paragraph()

    # ------------------------------------------------------------
    # Sections
    # ------------------------------------------------------------

    for section in content.sections:

        doc.add_heading(section.heading, level=1)

        for para_text in section.body.split("\n\n"):

            para_text = para_text.strip()

            if para_text:
                paragraph = doc.add_paragraph(para_text)
                paragraph.paragraph_format.space_after = Pt(10)

    # ------------------------------------------------------------
    # Temporary file
    # ------------------------------------------------------------

    tmp_dir = tempfile.gettempdir()

    filename = (
        f"{document_type.value}_"
        f"{datetime.now().strftime('%Y%m%d%H%M%S')}_"
        f"{uuid.uuid4().hex[:8]}.docx"
    )

    filepath = os.path.join(tmp_dir, filename)

    doc.save(filepath)

    return filepath


def _get_student_id_for_idea(idea_id: str) -> str:
    """
    Gets the internal student.id belonging to a project idea.
    """

    context = fetch_idea_with_student_context(idea_id)

    return context.idea.student_id


def store_generated_document(
    idea_id: str,
    document_type: DocumentType,
    filepath: str,
) -> dict:
    """
    Uploads the generated DOCX to Supabase Storage and records
    its metadata in the public.documents table.

    Storage structure:

    generated-documents/
        student_id/
            idea_id/
                document_type/
                    filename.docx
    """

    supabase = get_supabase()

    student_id = _get_student_id_for_idea(idea_id)

    document_id = str(uuid.uuid4())

    timestamp = datetime.now(timezone.utc).strftime(
        "%Y%m%d_%H%M%S"
    )

    filename = (
        f"{document_type.value}_"
        f"{timestamp}_"
        f"{document_id[:8]}.docx"
    )

    storage_path = (
        f"{student_id}/"
        f"{idea_id}/"
        f"{document_type.value}/"
        f"{filename}"
    )

    try:

        # --------------------------------------------------------
        # Read generated DOCX
        # --------------------------------------------------------

        with open(filepath, "rb") as file:
            file_bytes = file.read()

        # --------------------------------------------------------
        # Upload to private Supabase Storage
        # --------------------------------------------------------

        supabase.storage.from_(DOCUMENT_BUCKET).upload(
            storage_path,
            file_bytes,
            {
                "content-type": DOCX_MEDIA_TYPE,
                "cache-control": "3600",
                "upsert": False,
            },
        )

        # --------------------------------------------------------
        # Save metadata in public.documents
        # --------------------------------------------------------

        result = (
            supabase
            .table("documents")
            .insert(
                {
                    "id": document_id,
                    "student_id": student_id,
                    "idea_id": idea_id,
                    "document_type": document_type.value,
                    "file_name": filename,
                    "storage_path": storage_path,
                }
            )
            .execute()
        )

        # --------------------------------------------------------
        # If database insert failed, remove Storage file
        # --------------------------------------------------------

        if not result.data:

            try:
                supabase.storage.from_(DOCUMENT_BUCKET).remove(
                    [storage_path]
                )
            except Exception:
                pass

            raise RuntimeError(
                "The document was uploaded but its database "
                "record could not be created."
            )

        return result.data[0]

    finally:

        # --------------------------------------------------------
        # Remove temporary local copy
        # --------------------------------------------------------

        try:

            if os.path.exists(filepath):
                os.remove(filepath)

        except OSError:
            pass


def download_stored_document(storage_path: str) -> bytes:
    """
    Downloads a previously generated document from private
    Supabase Storage using the backend service-role client.
    """

    supabase = get_supabase()

    return supabase.storage.from_(DOCUMENT_BUCKET).download(
        storage_path
    )