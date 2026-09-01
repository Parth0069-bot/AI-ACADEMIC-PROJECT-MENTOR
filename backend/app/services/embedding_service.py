"""
Embedding pipeline for student project documents.

Flow:
    1. Get the document's raw text (either passed in directly, or
       pulled from Supabase Storage + extracted from the DOCX).
    2. Chunk the text into overlapping word-windows.
    3. Embed every chunk with a local sentence-transformers model.
    4. Store the chunks + vectors in the document_chunks table
       (pgvector) via embedding_repository.

The model is loaded once per process (module-level cache) since
loading it is the expensive part -- repeated requests reuse it.
"""

import io
import logging
from functools import lru_cache

from docx import Document

from app.core.config import settings
from app.core.document_repository import download_stored_document, get_latest_document
from app.core.embedding_repository import replace_chunks_for_document
from app.core.idea_repository import fetch_idea_with_student_context

logger = logging.getLogger(__name__)


@lru_cache
def get_embedding_model():
    """
    Lazily loads and caches the sentence-transformers model.

    Imported inside the function (rather than at module load time)
    so that importing this module doesn't pull in torch / the model
    download for code paths that never touch embeddings.
    """

    from sentence_transformers import SentenceTransformer

    logger.info("Loading embedding model: %s", settings.embedding_model)

    return SentenceTransformer(settings.embedding_model)


def chunk_text(
    text: str,
    chunk_size: int | None = None,
    overlap: int | None = None,
) -> list[str]:
    """
    Splits text into overlapping word-based chunks.

    Word-based (rather than character-based) chunking keeps chunks
    from cutting words in half, and roughly tracks token count well
    enough for a sentence-transformers model with a short context
    window. `overlap` repeats the tail of one chunk at the start of
    the next, so a sentence that happens to land on a chunk boundary
    still has surrounding context in at least one chunk.
    """

    chunk_size = chunk_size or settings.embedding_chunk_size
    overlap = overlap if overlap is not None else settings.embedding_chunk_overlap

    if overlap >= chunk_size:
        raise ValueError("embedding_chunk_overlap must be smaller than embedding_chunk_size")

    words = text.split()

    if not words:
        return []

    chunks = []
    step = chunk_size - overlap
    start = 0

    while start < len(words):
        window = words[start : start + chunk_size]
        chunk = " ".join(window).strip()

        if chunk:
            chunks.append(chunk)

        if start + chunk_size >= len(words):
            break

        start += step

    return chunks


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """
    Pulls plain text out of a DOCX file's paragraphs (headings,
    body text) and table cells. Good enough for the formal,
    prose-heavy documents this platform generates.
    """

    doc = Document(io.BytesIO(file_bytes))

    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text.strip())

    return "\n\n".join(parts)


def _resolve_document_text(idea_id: str, document_type: str) -> tuple[str, str | None]:
    """
    Fetches the latest stored DOCX for this idea + document type and
    extracts its text. Returns (text, document_id).
    """

    record = get_latest_document(idea_id=idea_id, document_type=document_type)

    if record is None:
        raise ValueError(
            f"No stored '{document_type}' document was found for idea {idea_id}. "
            "Generate the document first, or pass `text` directly."
        )

    file_bytes = download_stored_document(record["storage_path"])
    text = extract_text_from_docx_bytes(file_bytes)

    return text, record.get("id")


def process_document_for_embeddings(
    *,
    idea_id: str,
    document_type: str,
    text: str | None = None,
    document_id: str | None = None,
) -> dict:
    """
    Chunks, embeds, and stores a document's text in document_chunks.

    If `text` isn't supplied, the latest stored document of this
    type for `idea_id` is downloaded from Supabase Storage and its
    text is extracted automatically.
    """

    # Confirms the idea exists before doing any (comparatively
    # expensive) embedding work on it.
    fetch_idea_with_student_context(idea_id)

    if text is None:
        text, resolved_document_id = _resolve_document_text(idea_id, document_type)
        document_id = document_id or resolved_document_id

    raw_chunks = chunk_text(text)

    if not raw_chunks:
        return {
            "idea_id": idea_id,
            "document_id": document_id,
            "document_type": document_type,
            "chunks_stored": 0,
            "model": settings.embedding_model,
            "dimensions": settings.embedding_dimensions,
        }

    model = get_embedding_model()

    # normalize_embeddings=True makes cosine similarity (used by the
    # match_document_chunks RPC) equivalent to a plain dot product.
    vectors = model.encode(
        raw_chunks,
        normalize_embeddings=True,
        show_progress_bar=False,
    )

    chunk_rows = [
        {
            "chunk_index": index,
            "content": chunk,
            "token_count": len(chunk.split()),
            "embedding": vectors[index].tolist(),
        }
        for index, chunk in enumerate(raw_chunks)
    ]

    stored = replace_chunks_for_document(
        idea_id=idea_id,
        document_id=document_id,
        document_type=document_type,
        chunks=chunk_rows,
    )

    logger.info(
        "Embedded document idea=%s type=%s chunks=%d",
        idea_id,
        document_type,
        len(stored),
    )

    return {
        "idea_id": idea_id,
        "document_id": document_id,
        "document_type": document_type,
        "chunks_stored": len(stored),
        "model": settings.embedding_model,
        "dimensions": settings.embedding_dimensions,
    }


def embed_query(query: str) -> list[float]:
    """Embeds a single search query with the same model/normalization
    used for stored chunks, so cosine similarity is meaningful."""

    model = get_embedding_model()

    vector = model.encode(
        [query],
        normalize_embeddings=True,
        show_progress_bar=False,
    )[0]

    return vector.tolist()
